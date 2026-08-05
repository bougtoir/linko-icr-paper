#!/usr/bin/env python3
"""Build the LINKO manuscript files from the generated analysis results.

Outputs (in ``icr_paper/``):
* ``LINKO_manuscript_english.docx`` - Statistics in Medicine format
* ``LINKO_manuscript_japanese.docx`` - Japanese reference translation
* ``manuscript.md`` - plain-text mirror of the English manuscript

The text lives in :mod:`icr_paper.src.manuscript_content`; every number is read
from ``results/results.json``. Run ``python run_analysis.py`` first.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BASE = Path(__file__).resolve().parent
sys.path.insert(0, str(BASE.parent))

from icr_paper.src.figure_export import export_submission_figures
from icr_paper.src.manuscript_content import (
    build_english,
    build_japanese,
    figure_blocks,
    renumber_citations,
)
from icr_paper.src.omml_equations import build_omml
from icr_paper.src.results_loader import load_results

CITATION_RE = re.compile(r"\[(\d+(?:,\d+)*)\]")
EN_FONT = "Times New Roman"
JA_FONT = "Yu Mincho"


def _new_document(font: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    return doc


def _add_text(paragraph, text: str, font: str, italic: bool = False) -> None:
    """Add text, rendering ``[1,2]`` citations as superscript numerals."""
    position = 0
    for match in CITATION_RE.finditer(text):
        before = text[position : match.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.name = font
            run.italic = italic
        run = paragraph.add_run(match.group(1))
        run.font.name = font
        run.font.superscript = True
        position = match.end()
    tail = text[position:]
    if tail:
        run = paragraph.add_run(tail)
        run.font.name = font
        run.italic = italic


def _render_table(doc: Document, spec: dict, font: str) -> None:
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(12)
    run = caption.add_run(f"{spec['label']}. {spec['caption']}")
    run.bold = True
    run.font.name = font
    run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=len(spec["headers"]))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, spec["headers"]):
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(header))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = font
    for row_values in spec["rows"]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
            run.font.name = font


def _render_figure(doc: Document, spec: dict, font: str, width_inches: float = 6.0) -> None:
    path = Path(spec["path"])
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(12)
    run = caption.add_run(f"{spec['label']}. {spec['caption']}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = font


def render_docx(blocks: list, output: Path, font: str, embed_figures: bool) -> Path:
    doc = _new_document(font)
    figure_specs = []

    for kind, payload in blocks:
        if kind == "title":
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(payload)
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = font
        elif kind in ("h1", "h2", "h3"):
            heading = doc.add_heading(level=int(kind[1]))
            run = heading.add_run(payload)
            run.font.name = font
        elif kind == "p":
            paragraph = doc.add_paragraph()
            _add_text(paragraph, payload, font)
        elif kind == "eq":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.5)
            omml, trailing = build_omml(payload)
            paragraph._element.append(omml)
            if trailing:
                run = paragraph.add_run(trailing)
                run.font.name = font
        elif kind == "table":
            _render_table(doc, payload, font)
        elif kind == "figure":
            figure_specs.append(payload)
            if embed_figures:
                _render_figure(doc, payload, font)
        elif kind == "references":
            for index, reference in enumerate(payload, start=1):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(2)
                _add_text(paragraph, f"{index}. " + reference.replace("*", ""), font)
        elif kind == "pagebreak":
            doc.add_page_break()

    if figure_specs and not embed_figures:
        doc.add_page_break()
        heading = doc.add_heading(level=1)
        heading.add_run("Figure legends").font.name = font
        for spec in figure_specs:
            paragraph = doc.add_paragraph()
            _add_text(paragraph, f"{spec['label']}. {spec['caption']}", font)

    doc.save(str(output))
    return output


def render_tables_docx(blocks: list, output: Path, font: str) -> Path:
    """Separate editable file containing only the tables."""
    doc = _new_document(font)
    heading = doc.add_heading(level=1)
    heading.add_run("Tables").font.name = font
    for kind, payload in blocks:
        if kind == "table":
            _render_table(doc, payload, font)
            doc.add_paragraph()
    doc.save(str(output))
    return output


def render_markdown(blocks: list, output: Path) -> Path:
    lines = []
    for kind, payload in blocks:
        if kind == "title":
            lines.append(f"# {payload}\n")
        elif kind in ("h1", "h2", "h3"):
            lines.append("\n" + "#" * (int(kind[1]) + 1) + f" {payload}\n")
        elif kind == "p":
            lines.append(payload + "\n")
        elif kind == "eq":
            lines.append(f"\n    {payload}\n")
        elif kind == "table":
            lines.append(f"\n**{payload['label']}.** {payload['caption']}\n")
            lines.append("| " + " | ".join(payload["headers"]) + " |")
            lines.append("|" + "---|" * len(payload["headers"]))
            for row in payload["rows"]:
                lines.append("| " + " | ".join(str(v) for v in row) + " |")
            lines.append("")
        elif kind == "figure":
            rel = Path(payload["path"]).relative_to(BASE)
            lines.append(f"\n![{payload['label']}]({rel})\n")
            lines.append(f"*{payload['label']}. {payload['caption']}*\n")
        elif kind == "references":
            for index, reference in enumerate(payload, start=1):
                lines.append(f"{index}. {reference}")
            lines.append("")
    output.write_text("\n".join(lines))
    return output


def check_abstract_length(blocks: list, limit: int = 250) -> int:
    """Statistics in Medicine allows at most 250 words in the abstract."""
    for index, (kind, payload) in enumerate(blocks):
        if kind == "h1" and payload == "Abstract":
            text = blocks[index + 1][1]
            words = len(CITATION_RE.sub("", text).split())
            if words > limit:
                raise SystemExit(
                    f"Abstract is {words} words, exceeding the {limit}-word limit."
                )
            for position, (kind, payload) in enumerate(blocks):
                if kind == "p" and payload.startswith("Word count of abstract"):
                    blocks[position] = (
                        "p",
                        f"Word count of abstract: {words} (limit {limit}).",
                    )
            return words
    raise SystemExit("Abstract not found in manuscript blocks.")


def main() -> None:
    results = load_results()
    english = renumber_citations(build_english(results))
    japanese = build_japanese(results)

    words = check_abstract_length(english)
    print(f"Abstract: {words} words (limit 250)")

    # Statistics in Medicine: tables follow the references, figures are supplied
    # as separate files and represented in the manuscript by their legends.
    out_en = render_docx(
        english, BASE / "LINKO_manuscript_english.docx", EN_FONT, embed_figures=False
    )
    out_ja = render_docx(
        japanese, BASE / "LINKO_manuscript_japanese.docx", JA_FONT, embed_figures=False
    )
    # Reading copy with the figures placed after the paragraph that cites them.
    out_review = render_docx(
        english,
        BASE / "LINKO_manuscript_english_with_figures.docx",
        EN_FONT,
        embed_figures=True,
    )
    out_tables = render_tables_docx(
        english, BASE / "LINKO_tables_english.docx", EN_FONT
    )
    out_md = render_markdown(english, BASE / "manuscript.md")
    index = export_submission_figures(figure_blocks(english))
    for path in (out_en, out_ja, out_review, out_tables, out_md):
        print(f"Wrote {path}")
    print(f"Wrote {len(index)} separate figure files to figures/submission/")


if __name__ == "__main__":
    main()
